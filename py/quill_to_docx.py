"""从 Quill sqlite 数据库导出 docx 文件"""

import base64
import os
import re
import sqlite3
import json
import argparse
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_UNDERLINE, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from docx import Document


class Tools:
    """工具类"""

    @staticmethod
    def get_element(obj):
        """警告归一化"""
        return obj._element

    @staticmethod
    def get_tc(obj):
        """警告归一化"""
        return obj._tc

    @staticmethod
    def get_parent(obj):
        """警告归一化"""
        return obj._parent


# 定义字体映射表
class DocSys:
    """主要实现"""

    @staticmethod
    def add_hyperlink(run, url):
        """添加超链接"""
        part = run.part
        r_id = part.relate_to(
            url,
            reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        for child in Tools.get_element(run):
            hyperlink.append(child)
        Tools.get_element(run).clear()
        Tools.get_element(run).append(hyperlink)
        run.font.underline = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    string = {
        "w": "w:eastAsia",
        "temp_image":"temp_image.png"
    }

    FONT_MAP = {
        "kaiti": "KaiTi",  # 楷体
        "sans-serif": "Calibri",  # 无衬线字体（常用）
        "monospace": "Courier New",  # 等宽字体
        "simsun": "SimSun",  # 宋体
        "simhei": "SimHei",  # 黑体
        "fangsong": "FangSong",  # 仿宋
        "yahei": "Microsoft YaHei",  # 微软雅黑
        "default": "SimHei",
        # 可以继续添加其他映射
    }

    @staticmethod
    def export_docx_from_sqlite(db_path, output_path, generate_toc=False):
        """主要入口"""
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()

        doc = Document()

        if generate_toc:
            add_toc(doc)
            doc.add_page_break()

        # 获取按 sort_order 排序的所有章节
        cursor.execute("SELECT chapter_id, title FROM chapters ORDER BY sort_order")
        chapters = cursor.fetchall()

        for chapter_id, title in chapters:
            doc.add_heading(title, level=1)
            # 获取当前章节下的所有 page 内容
            cursor.execute(
                """SELECT content FROM pages 
                WHERE chapter_id = ? 
                ORDER BY page_num
            """,
                (chapter_id,),
            )
            pages = cursor.fetchall()
            for (content_blob,) in pages:
                if not content_blob:
                    continue
                try:
                    delta = json.loads(content_blob)
                    DocSys.write_delta_run(delta, doc)
                except json.JSONDecodeError as e:
                    print(f"解析 page 内容失败: {e}")
                    continue

        doc.save(output_path)
        conn.close()

    @staticmethod
    def preprocess_delta_ops_for_header(ops) -> list[dict]:
        """
        遍历 ops,将形如 {"attributes": {"header": 2}, "insert": "\n"} 的项，
        合并到它的前一个项中，作为其 attributes。
        """
        i = 0
        while i < len(ops):

            has_header = ops[i].get("attributes", {}).get("header", None)
            has_list = ops[i].get("attributes", {}).get("list", None)
            has_align = ops[i].get("attributes", {}).get("align", None)
            has_indent = ops[i].get("attributes", {}).get("indent", None)
            has_code_block = ops[i].get("attributes", {}).get("code-block", None)
            has_blockquote = ops[i].get("attributes", {}).get("blockquote", None)

            has_special = (
                has_header
                or has_list
                or has_align
                or has_indent
                or has_code_block
                or has_blockquote
            )

            if has_special:
                ops[i]["special"] = True
            i += 1
        return ops

    @staticmethod
    def write_delta_run(delta_ops, doc):
        """将 delta ops 写入 docx 的run段落"""
        current_paragraph = None
        preprocess_ops = DocSys.preprocess_delta_ops_for_header(
            delta_ops.get("ops", [])
        )

        def set_paragraph_style(paragraph, style):
            def convert_to_code_block(to_paragraph, background_color="fdf6e3"):
                """
                将给定的段落转换为具有指定背景颜色的代码块样式。
                :param to_paragraph: 需要转换为代码块样式的段落对象
                :param background_color: 十六进制颜色字符串，默认为淡橘色 (#FFDAB9)
                """
                doc_handle = Tools.get_parent(to_paragraph)
                original_text = to_paragraph.text
                to_paragraph.clear()

                # 在原段落位置插入一个单行单列的表格（注意这里加了 width 参数）
                table = doc_handle.add_table(
                    rows=1, cols=1, width=Cm(15)
                )  # 表格宽度设为 15 厘米
                cell = table.cell(0, 0)

                # 设置单元格背景色
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:fill"), background_color)  # 设置背景颜色
                Tools.get_tc(cell).get_or_add_tcPr().append(shading_elm)

                # 设置段落格式
                p_in_cell = cell.paragraphs[0]
                p_in_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # 添加原始段落的内容到表格单元格中，并设置等宽字体
                run = p_in_cell.add_run(original_text)
                run.font.name = DocSys.FONT_MAP["monospace"]
                Tools.get_element(run).rPr.rFonts.set(
                    qn(DocSys.string["w"]), DocSys.FONT_MAP["monospace"]
                )  # 中文兼容
                run.font.size = Pt(10)

                # 删除原始段落（现在它是空的）
                p = Tools.get_element(to_paragraph)
                p.getparent().remove(p)

            def convert_to_blockquote(
                paragraph, background_color="E6E6E6", left_border_color="4472C4"
            ):
                """
                将给定段落转换为 blockquote 引用块样式。
                :param paragraph: 需要转换的段落对象
                :param background_color: 背景颜色（HEX）
                :param left_border_color: 左侧边框颜色（HEX）
                """
                doc = Tools.get_parent(paragraph)
                original_text = paragraph.text
                paragraph.clear()

                # 插入表格（1行1列），宽度为页面宽度的 90%
                table = doc.add_table(rows=1, cols=1, width=Cm(15))
                cell = table.cell(0, 0)

                # 设置单元格背景色
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:fill"), background_color)
                Tools.get_tc(cell).get_or_add_tcPr().append(shading_elm)

                # 设置左侧边框线（可选）
                tc_borders = OxmlElement("w:tcBorders")
                left = OxmlElement("w:left")
                left.set(qn("w:val"), "single")
                left.set(qn("w:sz"), "8")  # 线宽
                left.set(qn("w:color"), left_border_color)
                tc_borders.append(left)
                Tools.get_tc(cell).get_or_add_tcPr().append(tc_borders)

                # 设置段落格式
                p_in_cell = cell.paragraphs[0]
                p_in_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # 添加原文内容，并设置引用样式
                run = p_in_cell.add_run(original_text)
                run.font.name = "Calibri"
                Tools.get_element(run).rPr.rFonts.set(qn(DocSys.string["w"]), "Calibri")
                run.font.size = Pt(11)
                run.font.italic = True

                # 删除原段落
                p = Tools.get_element(paragraph)
                p.getparent().remove(p)

            if "header" in style:
                level = int(style["header"])
                paragraph.style = "Heading " + str(level)
            elif "list" in style:
                list_type = style["list"]
                style_name = "List Bullet" if list_type == "bullet" else "List Number"
                paragraph.style = style_name
            elif "code-block" in style:
                convert_to_code_block(paragraph)
            elif "blockquote" in style:
                convert_to_blockquote(paragraph)
            if "align" in style:
                if style["align"] == "center":
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif style["align"] == "right":
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif style["align"] == "justify":
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if "indent" in style:
                paragraph.paragraph_format.left_indent = Pt(int(style["indent"]) * 30)

        def process_accumulated(text, paragraph, style):
            try:
                if text:
                    if paragraph is None:
                        paragraph = doc.add_paragraph()
                    run = paragraph.add_run(text)
                    DocSys.apply_attributes(run, style)
            except ValueError as e:
                print(e)
            return paragraph

        for op in preprocess_ops:
            if op.get("special", False):
                set_paragraph_style(current_paragraph, op["attributes"])
                current_paragraph = doc.add_paragraph()
            else:
                insert = op["insert"]
                if isinstance(insert, str):
                    lines = re.split("(\n)", insert)
                    for _, line in enumerate(lines):
                        if line == "\n":
                            current_paragraph = doc.add_paragraph()
                        else:
                            current_paragraph = process_accumulated(
                                line, current_paragraph, op.get("attributes", {})
                            )
                elif isinstance(insert, dict):  # 处理图片、视频等嵌入内容
                    if "image" in insert:
                        image_data = insert["image"]
                        if image_data.startswith("data:image"):
                            image_data = re.sub(
                                "^data:image/.+;base64,", "", image_data
                            )
                            image_bytes = base64.b64decode(image_data)
                            with open(DocSys.string["temp_image"], "wb") as f:
                                f.write(image_bytes)
                            new_paragraph = doc.add_paragraph()
                            run = new_paragraph.add_run()
                            run.add_picture(DocSys.string["temp_image"], width=Inches(2))
                            os.remove(DocSys.string["temp_image"])
                            current_paragraph = None
                    elif "video" in insert:
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run(f"[视频: {insert['video']}]")
                        run.italic = True
                        current_paragraph = None

    @staticmethod
    def apply_attributes(run, attrs):
        """应用属性到 run"""
        if "bold" in attrs and attrs["bold"]:
            run.bold = True
        if "italic" in attrs and attrs["italic"]:
            run.italic = True
        if "underline" in attrs:
            run.underline = (
                WD_UNDERLINE.SINGLE
                if attrs["underline"] is True
                else WD_UNDERLINE.DOUBLE
            )
        if "strike" in attrs and attrs["strike"]:
            run.font.strike = True
        if "color" in attrs:
            color = attrs["color"]
            if color.startswith("#"):
                color = color[1:]
            if len(color) == 6:
                r, g, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
                run.font.color.rgb = RGBColor(r, g, b)
        if "background" in attrs:
            bg_color = attrs["background"]
            if bg_color.startswith("#"):
                bg_color = bg_color[1:]
            try:
                # 构造一个带命名空间的 <w:shd> 节点
                shading_elm = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear" />'
                )
                # 清除已有的 shading
                rpr = Tools.get_element(run).get_or_add_rPr()
                for child in rpr:
                    if child.tag == f'{{{nsdecls("w")}}}shd':
                        rpr.remove(child)
                rpr.append(shading_elm)
            except (ValueError, SyntaxError) as e:
                print(f"[警告] 设置背景色失败: {e}")
        if "script" in attrs:
            if attrs["script"] == "sub":
                run.font.subscript = True
            elif attrs["script"] == "super":
                run.font.superscript = True
        if "font" in attrs:
            font_name = DocSys.FONT_MAP.get(attrs["font"].lower(), attrs["font"])
            # 设置西文字体
            run.font.name = font_name
            # 设置中文字体
            Tools.get_element(run).rPr.rFonts.set(qn(DocSys.string["w"]), font_name)
        else:
            run.font.name = DocSys.FONT_MAP["default"]
            Tools.get_element(run).rPr.rFonts.set(
                qn(DocSys.string["w"]), DocSys.FONT_MAP["default"]
            )
        if "size" in attrs:
            size_map = {"small": 8, "normal": 11, "large": 14, "huge": 18}
            size = size_map.get(attrs["size"], 11)
            run.font.size = Pt(size)
        if "link" in attrs:
            DocSys.add_hyperlink(run, attrs["link"])


# 自动目录占位符,第一次使用要更新一次
def add_toc(doc):
    """添加目录"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_string = "w:fldChar"  # 创建字段字符
    fld_char_type = "w:fldCharType"  # 创建字段字符类型
    fld_char = OxmlElement(fld_char_string)  # 创建字段字符
    fld_char.set(qn(fld_char_type), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'  # TOC 字段代码

    fld_sep = OxmlElement(fld_char_string)
    fld_sep.set(qn(fld_char_type), "separate")

    end_fld = OxmlElement(fld_char_string)
    end_fld.set(qn(fld_char_type), "end")

    Tools.get_element(run).append(fld_char)
    Tools.get_element(run).append(instr_text)
    Tools.get_element(run).append(fld_sep)
    Tools.get_element(run).append(end_fld)


def main():
    """ 命令行参数 """
    parser = argparse.ArgumentParser(description="Export DOCX from SQLite database.")
    parser.add_argument("db_path", type=str, help="Path to the SQLite database.")
    parser.add_argument(
        "output_path", type=str, help="Path to save the generated DOCX file."
    )
    parser.add_argument(
        "--generate-toc", action="store_true", help="Generate a table of contents."
    )

    args = parser.parse_args()

    # 调用导出函数
    DocSys.export_docx_from_sqlite(args.db_path, args.output_path, args.generate_toc)


if __name__ == "__main__":
    main()
